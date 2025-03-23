import os
import time
import mss
import keyboard
import pygetwindow as gw
from docx.enum.text import WD_ALIGN_PARAGRAPH
from plyer import notification
from tkinter import *
from tkinter import messagebox
from tkinter import filedialog
from tkinter import ttk
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Created by Jesus Peralta for Collections UAT #
# Constants:
FORBIDDEN_CHARACTERS = ['\\', '/', ':', '*', '?', '\"', '<', '>', '|']
DEFAULT_FILE_PATH = "./"
SAVE_PATH_FILE = 'save_path.txt'
USER = os.environ.get("USERNAME").title()
START_TIME = None
END_TIME = None


class Gui:
    def __init__(self):
        self.lines = None
        self.window = Tk()
        self.style = ttk.Style()
        self.style.theme_use('clam')
        self.window.title("SnapPy")
        self.window.config(padx=10, pady=10)
        self.doc_title_label = Label(text="Title:")
        self.doc_title_label.grid(column=0, row=0, sticky='w')
        self.doc_title = Entry(width=67)
        self.doc_title.grid(column=0, row=1, sticky='w', columnspan=2)
        self.account_eci_label = Label(text="Account/ECI:")
        self.account_eci_label.grid(column=0, row=2, sticky='w')
        self.account_number = Entry(width=67)
        self.account_number.insert(0, "(Optional)")
        self.account_number.grid(column=0, row=3, sticky='w', columnspan=2)
        self.gherkin_steps_label = Label(text="Gherkin Steps:")
        self.gherkin_steps_label.grid(column=0, row=4, sticky='w')
        self.default_file_path = "./"
        self.text = Text(height=30, width=50)
        self.text.insert(1.0, "Given...")
        self.text.grid(column=0, row=5, columnspan=2)

        self.generate = Button(text="Generate", command=self.check_fields)
        self.save_path = Button(text="Save Path", command=self.set_folder_path)
        self.generate.grid(column=0, row=6, pady=10)
        self.save_path.grid(column=1, row=6, pady=10)
        self.window.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.window.mainloop()

    def check_fields(self):
        word_title = self.doc_title.get()
        acct_number = self.account_number.get()

        if acct_number == '(Optional)':
            acct_number = 'N/A'

        if word_title == '':
            messagebox.showerror(
                title="Title Missing Error",
                message="Please ensure the title field is populated.")
        else:
            for word in word_title:
                if word in FORBIDDEN_CHARACTERS:
                    messagebox.showerror(
                        title="Invalid Character Error",
                        message="The below characters are not allowed in the title. Please rename the title.\n\n['\\','/',':','*','?','\"','<','>','|']")
                    break
            else:
                messagebox.showinfo(
                    title="Instructions",
                    message="To take a screenshot, please press 'left shift'. \nTo move to the next step, please press 'alt'. \nTo check what step you are on, please press '`'")
                self.get_text(word_title, acct_number)

    def get_text(self, word_title, acct_number):
        self.window.withdraw()
        self.doc = Document()
        self.sections = self.doc.sections
        save_path = self.get_save_path()
        START_TIME = datetime.now()
        for section in self.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = f"{word_title}\nAccount/ECI #: {acct_number}\n"
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.lines = self.text.get("1.0", "end-1c").splitlines()

        for i in range(len(self.lines)):
            line = self.lines[i]

            if i + 1 < len(self.lines):
                next_line = self.lines[i + 1]
            else:
                next_line = None

            self.doc.add_paragraph(line)
            self.take_screenshot(line, next_line)

        try:
            END_TIME = datetime.now()
            total_time = END_TIME - START_TIME
            hours, remainder = divmod(total_time.seconds, 3600)
            minutes, seconds = divmod(remainder, 60)

            self.doc.add_paragraph(f"Execution Time: {hours} Hours, {minutes} Minutes, {seconds} Seconds\nUser: {USER}")
            self.doc.save(f"{save_path}{word_title}.docx")
            self.window.deiconify()
            messagebox.showinfo(
                title="Information",
                message=f"Code has finished executing. The file has been saved to the below location\n\nFile path:\n{self.default_file_path}{word_title}.docx")
            self.delete_pictures()
        except PermissionError:
            self.doc.save(f"{self.default_file_path}{word_title}_v2.docx")
            self.window.deiconify()
            messagebox.showerror(
                title="Save Error",
                message=f"An issue occurred when saving the file.\n\nFile has been renamed to {word_title}_v2.docx\n\nPlease make sure file is closed before rerunning application")
        finally:
            self.lines.clear()
            self.text.delete("1.0", "end")

    def take_screenshot(self, message, next_message):
        initialized_message = message
        initialized_next_message = next_message
        while True:
            if keyboard.is_pressed('left shift'):
                try:
                    self.screenshot()
                    time.sleep(0.5)
                    self.show_notification(f"Captured for the below line:\n{message}")
                except AttributeError:
                    self.show_notification("Error: Please click the window you want to capture and try again.")
                    time.sleep(0.5)

            if keyboard.is_pressed('alt'):
                time.sleep(0.5)
                if next_message != None:
                    self.show_notification(f"You are on the below line:\n{next_message}")
                break

            if keyboard.is_pressed('`'):
                self.show_notification(f"You are currently on the below line:\n{message}")
                time.sleep(0.5)
                self.take_screenshot(initialized_message, initialized_next_message)
                break

    def screenshot(self):

        active_window = gw.getActiveWindow()

        if active_window:
            bbox = {
                "top": active_window.top,
                "left": active_window.left,
                "width": active_window.width,
                "height": active_window.height
            }

            with mss.mss() as sct:
                screenshot = sct.grab(bbox)  # Grab the area from all monitors
                filename = f"{self.default_file_path}screenshot_{time.time()}.png"
                mss.tools.to_png(screenshot.rgb, screenshot.size, output=filename)

        self.doc.add_picture(filename, width=Inches(7.5))
        self.doc.add_paragraph("\n")
        time.sleep(1)

    def set_folder_path(self):
        folder_selected = filedialog.askdirectory(title="Select folder to save files")
        if folder_selected:
            self.default_file_path = f"{folder_selected}/"
            with open(SAVE_PATH_FILE, 'w') as file:
                file.write(self.default_file_path)

    def delete_pictures(self):
        for filename in os.listdir(self.default_file_path):
            if filename.endswith('.png'):
                file_path = os.path.join(self.default_file_path, filename)
                try:
                    os.remove(file_path)
                except Exception as e:
                    pass

    def show_notification(self, message):
        notification.notify(title="SnapPy", message=message, timeout=1)

    def get_save_path(self):
        save_path_doc = './save_path.txt'
        if os.path.exists(save_path_doc):
            with open(save_path_doc, "r") as file:
                savelocation = file.readline()
        else:
            savelocation = './'

        return savelocation

    def on_closing(self):
        if messagebox.askokcancel("Exit SnapPy", "Do you want to exit the application?"):
            save_doc = './save_path.txt'
            if os.path.exists(save_doc):
                os.remove(save_doc)
            self.window.destroy()


if __name__ == "__main__":
    gui = Gui()

# Below is the command I use to make a pyinstaller .exe file
# pyinstaller --onefile --noconsole --hidden-import="plyer.platforms.win.notification" --hidden-import="plyer.platforms" SnapPy.py
